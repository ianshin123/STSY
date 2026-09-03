# -*- coding: utf-8 -*-
import re, io, os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SRC, OUT, BASE = sys.argv[1], sys.argv[2], os.path.dirname(os.path.abspath(sys.argv[1]))
KO = 'Noto Serif CJK KR'

doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(2.0)
USABLE = sec.page_width - sec.left_margin - sec.right_margin

st = doc.styles['Normal']
st.font.name = KO; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), KO)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.6

def setfont(run, size=None, bold=None, italic=None, color=None):
    run.font.name = KO
    run._element.rPr.rFonts.set(qn('w:eastAsia'), KO)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = color

TOK = re.compile(r'(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)', re.S)
def rich(par, text, size=None, base_italic=False, color=None):
    text = re.sub(r'</?sub>', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    for part in TOK.split(text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            setfont(par.add_run(part[2:-2]), size, True, base_italic, color)
        elif part.startswith('`') and part.endswith('`'):
            r = par.add_run(part[1:-1]); r.font.name='Consolas'
            if size: r.font.size = Pt(size)
            r.italic = base_italic
        elif part.startswith('*') and part.endswith('*'):
            setfont(par.add_run(part[1:-1]), size, None, True, color)
        else:
            setfont(par.add_run(part), size, None, base_italic, color)

lines = io.open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    ln = lines[i]; s = ln.strip()

    if not s or s == '---':
        i += 1; continue

    m = re.match(r'^(#{1,4})\s+(.*)$', s)
    if m:
        lvl, txt = len(m.group(1)), m.group(2)
        p = doc.add_paragraph()
        if lvl == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
            rich(p, txt, 16)
            for r in p.runs: r.bold = True
        else:
            p.paragraph_format.space_before = Pt(14 if lvl == 2 else 10)
            p.paragraph_format.space_after = Pt(4)
            rich(p, txt, {2: 12.5, 3: 11, 4: 10.5}[lvl])
            for r in p.runs: r.bold = True
        i += 1; continue

    if s.startswith('<sub>'):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        rich(p, s, 9.5)
        i += 1; continue

    m = re.match(r'^!\[[^\]]*\]\(([^)]+)\)$', s)
    if m:
        path = m.group(1)
        if not os.path.isabs(path): path = os.path.join(BASE, path)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        try: p.add_run().add_picture(path, width=USABLE)
        except Exception as e: rich(p, '[그림 없음: %s]' % os.path.basename(path), 9)
        i += 1; continue

    if s.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            r = lines[i].strip()
            if not re.match(r'^\|[\s:\-\|]+\|$', r):
                rows.append([c.strip() for c in r.strip('|').split('|')])
            i += 1
        if rows:
            n = max(len(r) for r in rows)
            t = doc.add_table(rows=len(rows), cols=n)
            t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(n):
                    cell = t.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = 1.15
                    rich(p, row[ci] if ci < len(row) else '', 9.5)
                    if ri == 0:
                        for r in p.runs: r.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue

    if s.startswith('>'):
        buf = []
        while i < len(lines) and lines[i].strip().startswith('>'):
            buf.append(re.sub(r'^>\s?', '', lines[i].strip())); i += 1
        groups, cur = [], []
        for b in buf:
            b = b.strip()
            hm = re.match(r'^(#{1,4})\s+(.*)$', b)
            bullet = re.match(r'^[-*]\s+', b)
            if not b or hm or bullet:
                if cur: groups.append(('p', ' '.join(cur))); cur = []
                if hm: groups.append(('h', hm.group(2)))
                elif bullet: groups.append(('b', re.sub(r'^[-*]\s+', '', b)))
            else:
                cur.append(b)
        if cur: groups.append(('p', ' '.join(cur)))
        for kind, txt in groups:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 if kind == 'b' else 0.6)
            p.paragraph_format.space_after = Pt(2)
            if kind == 'h':
                rich(p, txt, 11)
                for r in p.runs: r.bold = True
            elif kind == 'b':
                rich(p, '· ' + txt, 10)
            else:
                rich(p, txt, 10)
        continue

    m = re.match(r'^(\s*)[-*]\s+(.*)$', ln)
    if m:
        depth = 1 if len(m.group(1)) >= 2 else 0
        buf = m.group(2)
        while i + 1 < len(lines) and re.match(r'^\s{2,}\S', lines[i+1]) and not re.match(r'^\s*[-*0-9]', lines[i+1].strip()):
            i += 1; buf += ' ' + lines[i].strip()
        p = doc.add_paragraph(style='List Bullet' if depth == 0 else 'List Bullet 2')
        p.paragraph_format.space_after = Pt(2)
        rich(p, buf)
        i += 1; continue

    m = re.match(r'^(\d+)\.\s+(.*)$', s)
    if m:
        buf = m.group(2)
        while i + 1 < len(lines) and re.match(r'^\s{2,}\S', lines[i+1]) and not re.match(r'^\s*\d+\.', lines[i+1].strip()):
            i += 1; buf += ' ' + lines[i].strip()
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        rich(p, buf)
        i += 1; continue

    # caption block (starts with *그림)
    if s.startswith('*그림'):
        buf = []
        while i < len(lines) and lines[i].strip():
            buf.append(lines[i].strip()); i += 1
        txt = ' '.join(buf).strip('*')
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        rich(p, txt, 9.5, base_italic=True)
        continue

    buf = [s]
    while i + 1 < len(lines) and lines[i+1].strip() and not re.match(r'^(#|\||>|\s*[-*]\s|\d+\.\s|!\[|<sub>|---$)', lines[i+1].strip()):
        i += 1; buf.append(lines[i].strip())
    p = doc.add_paragraph(); rich(p, ' '.join(buf))
    i += 1

doc.save(OUT)
print('saved', OUT)
